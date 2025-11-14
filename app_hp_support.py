# -*- coding: utf-8 -*-
from flask import Flask, render_template, jsonify, request, send_file, Response
from flask_cors import CORS
import os
import json
from datetime import datetime
import uuid
import traceback
import requests
import io
import google.generativeai as genai
from google.cloud import storage
from docx import Document
from bs4 import BeautifulSoup
from prompt_manager import PromptManager
import logging
# ★★★ 追加部分 1: 必要なライブラリをインポート ★★★
from google.cloud import texttospeech
import base64

# ロギング設定
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# 環境変数
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyBpFoecz-RcuZrVDKlrkOhBq1sDyGcnteQ")
ASTRO_BUILD_SERVICE_URL = os.getenv("ASTRO_BUILD_SERVICE_URL", "https://astro-builder-v2-393419998935.us-central1.run.app")
GCS_OUTPUT_BUCKET = os.getenv("GCS_OUTPUT_BUCKET", "gen-lang-client-0691275473-preview")
GCS_OUTPUT_PATH = os.getenv("GCS_OUTPUT_PATH", "current")
DEFAULT_PREVIEW_URL = os.getenv("DEFAULT_PREVIEW_URL", "https://gen-lang-client-0691275473.web.app")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

app = Flask(__name__)
CORS(app)

class AppState:
    def __init__(self):
        self.sessions = {}
        self.gemini_model = None

state = AppState()

# プロンプトマネージャーを初期化
prompt_manager = PromptManager()

# ★★★ 追加部分 2: TTSクライアントのグローバル変数と初期化関数 ★★★
tts_client = None

def initialize_tts_client():
    """Google Cloud TTS クライアントを初期化"""
    global tts_client
    try:
        tts_client = texttospeech.TextToSpeechClient()
        logger.info("Google Cloud TTS client initialized")
        return True
    except Exception as e:
        logger.error(f"Failed to initialize TTS client: {e}")
        return False

# ★★★ 修正部分 1: アプリ起動時に各種サービスを初期化 ★★★
@app.before_request
def initialize_services():
    """アプリ起動時にプロンプトとTTSクライアントを一度だけ初期化する"""
    if not hasattr(app, 'services_initialized'):
        prompt_manager.get("fix_instructions")  # 初回アクセスで自動読み込み
        initialize_tts_client() # TTSクライアントを初期化
        app.services_initialized = True
        logger.info("All services initialized")

@app.route("/")
def index():
    return render_template("index.html")

# ★★★ 修正部分 2: ヘルスチェックにTTSの状態を追加 ★★★
@app.route("/health")
def health():
    return jsonify({
        "status": "healthy",
        "service": "hp-support",
        "gemini_configured": bool(GEMINI_API_KEY),
        "tts_configured": tts_client is not None,  # TTSが初期化されているかを確認
        "prompts_bucket": bool(os.getenv("PROMPTS_BUCKET_NAME")),
        "preview_bucket": f"{GCS_OUTPUT_BUCKET}/{GCS_OUTPUT_PATH}",
        "default_preview_url": DEFAULT_PREVIEW_URL
    })

# プレビュー配信用エンドポイント(全アセット対応フォールバック付き)
@app.route("/preview")
@app.route("/preview/")
@app.route("/preview/<path:filename>")
def serve_preview(filename="index.html"):
    """GCSからビルド済みファイルを取得して配信(全アセット対応フォールバック付き)"""

    # Content-Type判定(早期に実行)
    content_type = 'text/html'
    if filename.endswith('.css'):
        content_type = 'text/css'
    elif filename.endswith('.js'):
        content_type = 'application/javascript'
    elif filename.endswith('.json'):
        content_type = 'application/json'
    elif filename.endswith('.png'):
        content_type = 'image/png'
    elif filename.endswith('.jpg') or filename.endswith('.jpeg'):
        content_type = 'image/jpeg'
    elif filename.endswith('.svg'):
        content_type = 'image/svg+xml'
    elif filename.endswith('.woff') or filename.endswith('.woff2'):
        content_type = 'font/woff2'
    elif filename.endswith('.ttf'):
        content_type = 'font/ttf'
    elif filename.endswith('.ico'):
        content_type = 'image/x-icon'
    elif filename.endswith('.webp'):
        content_type = 'image/webp'
    elif filename.endswith('.gif'):
        content_type = 'image/gif'

    # GCS接続を試みる（失敗した場合は即座にフォールバック）
    gcs_available = False
    storage_client = None
    bucket = None
    blob = None

    try:
        storage_client = storage.Client()
        bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)
        blob_path = f"{GCS_OUTPUT_PATH}/{filename}"
        blob = bucket.blob(blob_path)
        gcs_available = True

        # GCSにファイルがある場合は取得して返す
        if blob.exists():
            try:
                content = blob.download_as_bytes()
                logger.info(f"✓ Served from GCS: {blob_path} ({len(content)} bytes)")

                # index.htmlの場合、ベースURLと選択検知スクリプトを注入
                if filename == "index.html" and content_type == 'text/html':
                    content = inject_scripts_to_html(content)

                return Response(content, mimetype=content_type)
            except Exception as download_error:
                logger.warning(f"Failed to download from GCS: {download_error}, falling back")
                gcs_available = False
        else:
            logger.info(f"Preview file not found in GCS: {blob_path}, falling back to default URL")

    except Exception as gcs_error:
        # GCS接続エラー（認証失敗、ネットワークエラーなど）
        logger.warning(f"GCS access failed for {filename}: {gcs_error}, falling back to default URL")
        gcs_available = False

    # GCSが使えない場合、または、ファイルが見つからない場合はDEFAULT_PREVIEW_URLからフェッチ
    try:
        # DEFAULT_PREVIEW_URL からフォールバック
        # index.htmlの場合はルートURL、それ以外はパスを維持
        if filename == "index.html":
            fallback_url = DEFAULT_PREVIEW_URL
        else:
            fallback_url = f"{DEFAULT_PREVIEW_URL}/{filename}"

        logger.info(f"Fetching from fallback URL: {fallback_url}")
        fallback_response = requests.get(fallback_url, timeout=10, allow_redirects=True)

        if fallback_response.status_code == 200:
            content = fallback_response.content
            # レスポンスのContent-Typeを優先
            response_content_type = fallback_response.headers.get('Content-Type', content_type)
            # Content-Typeからcharset等を除去してシンプルに
            if ';' in response_content_type:
                response_content_type = response_content_type.split(';')[0].strip()

            logger.info(f"✓ Served from fallback URL: {fallback_url} ({len(content)} bytes, {response_content_type})")

            # index.htmlの場合、ベースURLと選択検知スクリプトを注入
            if filename == "index.html" and 'text/html' in response_content_type:
                content = inject_scripts_to_html(content)

            # GCSにキャッシュ保存を試みる（失敗しても続行）
            if gcs_available and blob is not None:
                try:
                    blob.upload_from_string(content, content_type=response_content_type)
                    logger.info(f"✓ Cached to GCS: {blob_path}")
                except Exception as cache_error:
                    logger.warning(f"Failed to cache to GCS (non-critical): {cache_error}")

            return Response(content, mimetype=response_content_type)
        else:
            logger.warning(f"Fallback URL returned {fallback_response.status_code} for {fallback_url}")
            return f"Preview file not found: {filename}", 404

    except Exception as fallback_error:
        logger.error(f"Fallback fetch failed for {filename}: {fallback_error}")
        traceback.print_exc()
        return f"Error serving preview file: {filename}", 500


def inject_scripts_to_html(content):
    """HTMLコンテンツにベースURLと選択検知スクリプトを注入する"""
    try:
        html_content = content.decode('utf-8')

        # <head>タグの直後に<base>タグを挿入(既にない場合のみ)
        if '<head>' in html_content and '<base' not in html_content:
            html_content = html_content.replace(
                '<head>',
                '<head>\n    <base href="/preview/">'
            )
            logger.info("✓ Injected <base href='/preview/'> tag")

        # 選択検知スクリプトを注入(既にない場合のみ)
        selection_script = """
<script>
// 親ウィンドウに選択情報を送信
document.addEventListener('mouseup', function() {
    setTimeout(function() {
        const selection = window.getSelection();
        const text = selection.toString().trim();
        if (text && text.length > 0) {
            const range = selection.getRangeAt(0);
            const container = range.commonAncestorContainer;
            const element = container.nodeType === 3 ? container.parentElement : container;

            // セレクタを生成
            let selector = element.tagName.toLowerCase();
            if (element.id) selector += '#' + element.id;
            if (element.className) selector += '.' + element.className.split(' ').join('.');

            window.parent.postMessage({
                type: 'text-selected',
                text: text,
                tagName: element.tagName,
                className: element.className,
                id: element.id,
                selector: selector
            }, '*');

            console.log('[IFRAME] Selection sent to parent:', text);
        }
    }, 10);
});
</script>
"""

        if '</body>' in html_content and 'text-selected' not in html_content:
            html_content = html_content.replace('</body>', selection_script + '</body>')
            logger.info("✓ Injected selection detection script")
        elif '</html>' in html_content and 'text-selected' not in html_content:
            # </body>がない場合は</html>の前に挿入
            html_content = html_content.replace('</html>', selection_script + '</html>')
            logger.info("✓ Injected selection detection script (before </html>)")

        return html_content.encode('utf-8')

    except Exception as e:
        logger.warning(f"Failed to inject scripts: {e}")
        return content

@app.route("/api/sessions", methods=["POST"])
def create_session():
    data = request.json
    session_id = str(uuid.uuid4())
    state.sessions[session_id] = {
        "id": session_id,
        "created_at": datetime.now().isoformat(),
        "case_type": data.get("caseType", "new"),
        "client_info": data.get("clientInfo", {}),
        "conversation_log": [],
        "build_jobs": [],
        "fix_instructions": []
    }
    return jsonify({"success": True, "sessionId": session_id}), 201

@app.route("/api/sessions/<session_id>", methods=["GET"])
def get_session(session_id):
    session = state.sessions.get(session_id)
    return jsonify({"success": True, "session": session}) if session else (jsonify({"error": "Session not found"}), 404)

@app.route("/chat-message", methods=["POST"])
def chat_message():
    data = request.json
    user_text = data.get("message", "").strip()
    session_id = data.get("session_id")
    
    if not user_text:
        return jsonify({"ai_response": "メッセージが空です"})
    
    if not GEMINI_API_KEY:
        return jsonify({"ai_response": "Gemini APIキーが設定されていません"})
    
    try:
        if not state.gemini_model:
            state.gemini_model = genai.GenerativeModel("gemini-2.0-flash-exp")
        
        # システムプロンプトを取得
        system_prompt = prompt_manager.get("chat_system")
        
        # システムプロンプト + ユーザーメッセージ
        full_prompt = f"{system_prompt}\n\nユーザー: {user_text}"
        
        response = state.gemini_model.generate_content(full_prompt)
        ai_response = response.text if hasattr(response, 'text') else "応答を生成できませんでした"
        
        if session_id and session_id in state.sessions:
            state.sessions[session_id]["conversation_log"].append({
                "timestamp": datetime.now().isoformat(),
                "type": "chat",
                "user": user_text,
                "ai": ai_response
            })
        
        return jsonify({"ai_response": ai_response})
    except Exception as e:
        logger.error(f"Chat error: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/api/trigger-build", methods=["POST"])
def trigger_build():
    data = request.json
    session_id = data.get("session_id")
    
    if not session_id or session_id not in state.sessions:
        return jsonify({"success": False, "error": "有効なセッションIDが必要です"}), 400
    
    try:
        build_response = requests.post(
            f"{ASTRO_BUILD_SERVICE_URL}/build",
            json={"session_id": session_id, "diffData": data.get("diffData", {})},
            timeout=30
        )
        
        if build_response.ok:
            build_data = build_response.json()
            if build_data.get("success") and "jobId" in build_data:
                state.sessions[session_id]["build_jobs"].append({
                    "job_id": build_data["jobId"],
                    "triggered_at": datetime.now().isoformat(),
                    "status": build_data.get("status", "pending")
                })
            return jsonify(build_data)
        else:
            return jsonify({"success": False, "error": f"HTTP {build_response.status_code}"}), build_response.status_code
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/build-status/<job_id>", methods=["GET"])
def get_build_status(job_id):
    try:
        status_response = requests.get(f"{ASTRO_BUILD_SERVICE_URL}/build/{job_id}", timeout=10)
        return jsonify(status_response.json()) if status_response.ok else (jsonify({"success": False, "error": f"HTTP {status_response.status_code}"}), status_response.status_code)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/add-selection", methods=["POST"])
def add_selection():
    data = request.json
    session_id = data.get("session_id")
    selection = data.get("selection")
    
    if not session_id or not selection:
        return jsonify({"success": False, "error": "session_idとselectionが必要です"}), 400
    
    session = state.sessions.get(session_id)
    if not session:
        return jsonify({"error": "Session not found"}), 404
    
    session["conversation_log"].append({
        "timestamp": datetime.now().isoformat(),
        "type": "selection",
        "data": selection,
        "user_comment": ""
    })
    
    # プロンプトマネージャーから質問文を生成
    auto_question = prompt_manager.get(
        "selection_analysis",
        selection_content=selection.get('content', ''),
        selection_type=selection.get('type', ''),
        user_comment=""
    )
    
    return jsonify({
        "success": True,
        "selection_id": len(session["conversation_log"]) - 1,
        "auto_question": auto_question
    })

@app.route("/api/upload-file", methods=["POST"])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "ファイルがありません"}), 400
    
    file = request.files['file']
    session_id = request.form.get('session_id')
    
    if not session_id or session_id not in state.sessions:
        return jsonify({"success": False, "error": "有効なセッションIDが必要です"}), 400
    
    file_url = f"https://storage.googleapis.com/your-bucket/{session_id}/{file.filename}"
    
    state.sessions[session_id]["conversation_log"].append({
        "timestamp": datetime.now().isoformat(),
        "type": "file_upload",
        "data": {
            "filename": file.filename,
            "purpose": request.form.get('purpose', ''),
            "url": file_url
        }
    })
    
    return jsonify({"success": True, "file_url": file_url})

@app.route("/api/generate-fix-instructions", methods=["POST"])
def generate_fix_instructions():
    data = request.json
    session_id = data.get("session_id")
    
    session = state.sessions.get(session_id)
    if not session or not session.get("conversation_log"):
        return jsonify({"success": False, "error": "セッションまたは会話ログがありません"}), 400
    
    if not GEMINI_API_KEY:
        return jsonify({"success": False, "error": "Gemini APIキーが設定されていません"}), 500
    
    try:
        if not state.gemini_model:
            state.gemini_model = genai.GenerativeModel("gemini-2.0-flash-exp")
        
        # 会話ログを整形
        conversation_text = "\n".join([
            f"[{log.get('timestamp', '')}] ユーザー: {log.get('user', '')} / AI: {log.get('ai', '')}"
            for log in session["conversation_log"]
            if log.get("type") == "chat"
        ])
        
        # プロンプトマネージャーからプロンプトを取得
        prompt = prompt_manager.get(
            "fix_instructions",
            timestamp=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            session_id=session_id,
            conversation_text=conversation_text
        )
        
        logger.info(f"Generating fix instructions with {len(conversation_text)} chars of conversation")
        
        response = state.gemini_model.generate_content(prompt)
        fix_instructions = response.text if hasattr(response, 'text') else "生成に失敗しました"
        
        session.setdefault("fix_instructions", []).append({
            "instructions": fix_instructions,
            "generated_at": datetime.now().isoformat()
        })
        
        return jsonify({"success": True, "fix_instructions": fix_instructions})
    except Exception as e:
        logger.error(f"Fix instructions generation error: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/download-word", methods=["POST"])
def download_word():
    data = request.json
    session_id = data.get("session_id")
    
    session = state.sessions.get(session_id)
    if not session or not session.get("fix_instructions"):
        return jsonify({"success": False, "error": "修正指示書がありません"}), 400
    
    fix_instructions = session["fix_instructions"][-1]["instructions"]
    
    doc = Document()
    doc.add_heading('修正指示書', 0)
    
    soup = BeautifulSoup(fix_instructions, 'html.parser')
    for element in soup.find_all(['h1', 'h2', 'h3', 'p']):
        if element.name.startswith('h'):
            doc.add_heading(element.get_text(), level=int(element.name[1]))
        else:
            doc.add_paragraph(element.get_text())
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'fix_instructions_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    )

# 管理API:プロンプト再読み込み(テスト用)
@app.route("/api/admin/reload-prompts", methods=["POST"])
def reload_prompts():
    """プロンプトを強制再読み込み(開発・テスト用)"""
    try:
        result = prompt_manager.reload()
        return jsonify({
            "success": True,
            "message": "プロンプトを再読み込みしました",
            **result
        })
    except Exception as e:
        logger.error(f"Prompt reload error: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

# 管理API:現在のプロンプト一覧(テスト用)
@app.route("/api/admin/prompts", methods=["GET"])
def list_prompts():
    """現在読み込まれているプロンプトの一覧を返す"""
    try:
        prompts_info = {
            name: {
                "length": len(content),
                "preview": content[:100] + "..." if len(content) > 100 else content
            }
            for name, content in prompt_manager.prompts.items()
        }
        
        return jsonify({
            "success": True,
            "prompts": prompts_info,
            "last_loaded": prompt_manager.last_loaded.isoformat() if prompt_manager.last_loaded else None,
            "cache_minutes": prompt_manager.cache_minutes
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# ★★★ 追加部分 3: 音声合成APIエンドポイント ★★★
@app.route("/api/tts/synthesize", methods=["POST"])
def synthesize_speech():
    """テキストを音声に変換して返す"""
    if not tts_client:
        return jsonify({"success": False, "error": "TTS client not initialized"}), 500
    
    data = request.json
    text = data.get("text", "").strip()
    
    if not text:
        return jsonify({"success": False, "error": "テキストが空です"}), 400
    
    # テキストが長すぎる場合は切り詰める (Google Cloud TTS の制限: 5000文字)
    if len(text) > 5000:
        text = text[:5000]
        logger.warning(f"Text truncated to 5000 characters")
    
    try:
        # 音声合成リクエストの設定
        synthesis_input = texttospeech.SynthesisInput(text=text)
        
        # 音声パラメータの設定
        voice = texttospeech.VoiceSelectionParams(
            language_code=data.get("language_code", "ja-JP"),
            name=data.get("voice_name", "ja-JP-Neural2-B"),  # 自然な女性の声
            ssml_gender=texttospeech.SsmlVoiceGender.FEMALE
        )
        
        # オーディオ設定
        audio_config = texttospeech.AudioConfig(
            audio_encoding=texttospeech.AudioEncoding.MP3,
            speaking_rate=data.get("speaking_rate", 1.0),  # 0.25 ~ 4.0
            pitch=data.get("pitch", 0.0),  # -20.0 ~ 20.0
            volume_gain_db=data.get("volume_gain_db", 0.0)  # -96.0 ~ 16.0
        )
        
        # 音声合成実行
        response = tts_client.synthesize_speech(
            input=synthesis_input,
            voice=voice,
            audio_config=audio_config
        )
        
        # Base64エンコードして返す
        audio_base64 = base64.b64encode(response.audio_content).decode('utf-8')
        
        logger.info(f"TTS synthesized: {len(text)} chars -> {len(response.audio_content)} bytes")
        
        return jsonify({
            "success": True,
            "audio": audio_base64,
            "format": "mp3",
            "text_length": len(text)
        })
        
    except Exception as e:
        logger.error(f"TTS synthesis error: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

# ★★★ 追加部分 4: 利用可能な音声一覧を取得するAPIエンドポイント ★★★
@app.route("/api/tts/voices", methods=["GET"])
def list_voices():
    """利用可能な音声一覧を取得"""
    if not tts_client:
        return jsonify({"success": False, "error": "TTS client not initialized"}), 500
    
    try:
        language_code = request.args.get("language_code", "ja-JP")
        
        # 音声一覧を取得
        response = tts_client.list_voices(language_code=language_code)
        
        voices = []
        for voice in response.voices:
            voices.append({
                "name": voice.name,
                "language_codes": voice.language_codes,
                "ssml_gender": texttospeech.SsmlVoiceGender(voice.ssml_gender).name,
                "natural_sample_rate_hertz": voice.natural_sample_rate_hertz
            })
        
        return jsonify({
            "success": True,
            "voices": voices,
            "count": len(voices)
        })
        
    except Exception as e:
        logger.error(f"List voices error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


# ========================================
# ★★★ immediate修正機能追加 ★★★
# 以下3つのエンドポイントを追加
# ========================================

# ★★★ 新規追加 1: /api/chat エンドポイント ★★★
@app.route("/api/chat", methods=["POST"])
def chat():
    """immediate/batch判定と修正指示生成"""
    try:
        data = request.json
        message = data.get("message", "")
        selection = data.get("selection")
        history = data.get("history", [])
        session_id = data.get("session_id")
        
        if not message:
            return jsonify({"success": False, "error": "メッセージが空です"}), 400
        
        # プロンプト取得
        system_prompt = prompt_manager.get("chat_system")
        
        # Gemini API呼び出し
        model = genai.GenerativeModel("gemini-2.0-flash-exp")
        
        user_prompt = f"""
選択情報:
{json.dumps(selection, ensure_ascii=False, indent=2) if selection else "なし"}

会話履歴:
{json.dumps(history[-3:], ensure_ascii=False, indent=2) if history else "なし"}

ユーザーメッセージ:
{message}

上記を分析し、JSON形式で返答してください。
"""
        
        response = model.generate_content(
            [system_prompt, user_prompt],
            generation_config={
                "temperature": 0.7,
                "response_mime_type": "application/json"
            }
        )
        
        # レスポンスをパース
        result = json.loads(response.text)
        
        # セッションに会話ログを記録
        if session_id and session_id in state.sessions:
            session = state.sessions[session_id]
            session["conversation_log"].append({
                "timestamp": datetime.now().isoformat(),
                "type": "modification_chat",
                "user": message,
                "assistant": result.get("response", ""),
                "action": result.get("action"),
                "modification": result.get("modification"),
                "selection": selection
            })
        
        return jsonify({
            "success": True,
            "action": result.get("action", "question"),
            "response": result.get("response", ""),
            "modification": result.get("modification")
        })
    
    except Exception as e:
        logger.error(f"チャットエラー: {e}")
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

# ★★★ 新規追加 2: /api/save-html エンドポイント ★★★
@app.route("/api/save-html", methods=["POST"])
def save_html():
    """修正後のHTMLを保存"""
    try:
        data = request.json
        html = data.get("html", "")
        modifications = data.get("modifications", [])
        
        if not html:
            return jsonify({"success": False, "error": "HTMLが空です"}), 400
        
        # GCS設定
        gcs_bucket = os.getenv("GCS_BUCKET_NAME", "ai-meeting-cloud.appspot.com")
        
        # GCSに保存
        storage_client = storage.Client()
        bucket = storage_client.bucket(gcs_bucket)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        blob_name = f"modified_html/index_{timestamp}.html"
        
        blob = bucket.blob(blob_name)
        blob.upload_from_string(html, content_type="text/html")
        
        logger.info(f"Modified HTML saved: {blob_name}")
        
        # 修正ログも保存
        log_blob_name = f"modification_logs/log_{timestamp}.json"
        log_blob = bucket.blob(log_blob_name)
        log_blob.upload_from_string(
            json.dumps({
                "timestamp": timestamp,
                "modifications": modifications
            }, ensure_ascii=False, indent=2),
            content_type="application/json"
        )
        
        logger.info(f"Modification log saved: {log_blob_name}")
        
        return jsonify({
            "success": True,
            "saved_path": blob_name,
            "log_path": log_blob_name
        })
    
    except Exception as e:
        logger.error(f"HTML保存エラー: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

# ★★★ 新規追加 3: /api/import-site エンドポイント ★★★
@app.route("/api/import-site", methods=["POST"])
def import_site():
    """外部サイトのHTMLをインポート"""
    try:
        data = request.json
        url = data.get("url", "https://eentry.co.jp")
        
        logger.info(f"Importing site: {url}")
        
        # URLからHTMLを取得
        response = requests.get(url, timeout=10, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        })
        
        if response.status_code == 200:
            html_content = response.text
            
            # 選択検知スクリプトを埋め込む
            selection_script = """
<script>
// 親ウィンドウに選択情報を送信
document.addEventListener('mouseup', function() {
    setTimeout(function() {
        const selection = window.getSelection();
        const text = selection.toString().trim();
        if (text && text.length > 0) {
            const range = selection.getRangeAt(0);
            const container = range.commonAncestorContainer;
            const element = container.nodeType === 3 ? container.parentElement : container;
            
            // セレクタを生成
            let selector = element.tagName.toLowerCase();
            if (element.id) selector += '#' + element.id;
            if (element.className) selector += '.' + element.className.split(' ').join('.');
            
            window.parent.postMessage({
                type: 'text-selected',
                text: text,
                tagName: element.tagName,
                className: element.className,
                id: element.id,
                selector: selector
            }, '*');
            
            console.log('[IFRAME] Selection sent to parent:', text);
        }
    }, 10);
});
</script>
"""
            # </body>の直前にスクリプトを挿入
            if '</body>' in html_content:
                html_content = html_content.replace('</body>', selection_script + '</body>')
            else:
                # </body>がない場合は末尾に追加
                html_content += selection_script
            
            # GCS設定
            gcs_bucket = os.getenv("GCS_BUCKET_NAME", "ai-meeting-cloud.appspot.com")
            
            # GCSに保存
            storage_client = storage.Client()
            bucket = storage_client.bucket(gcs_bucket)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            blob_name = f"imported_sites/site_{timestamp}.html"
            
            blob = bucket.blob(blob_name)
            blob.upload_from_string(html_content, content_type="text/html")
            
            # 公開URL生成
            blob.make_public()
            public_url = blob.public_url
            
            logger.info(f"Site imported and saved: {blob_name}")
            
            return jsonify({
                "success": True,
                "preview_url": public_url,
                "original_url": url,
                "saved_path": blob_name,
                "message": "サイトを読み込みました"
            })
        else:
            return jsonify({
                "success": False,
                "error": f"サイトの取得に失敗: HTTP {response.status_code}"
            }), 400
    
    except Exception as e:
        logger.error(f"サイトインポートエラー: {e}")
        traceback.print_exc()
        return jsonify({
            "success": False,
            "error": str(e)
        }), 500

# ★★★ 新規追加: /static/ 配下のファイル配信エンドポイント ★★★
@app.route("/static/<path:filename>")
def serve_static(filename):
    """staticディレクトリ配下のファイルを配信"""
    try:
        static_dir = os.path.join(os.path.dirname(__file__), 'static')
        file_path = os.path.join(static_dir, filename)

        if not os.path.exists(file_path):
            logger.error(f"Static file not found: {file_path}")
            return f"File not found: {filename}", 404

        # Content-Type判定
        content_type = 'application/octet-stream'
        if filename.endswith('.js'):
            content_type = 'application/javascript'
        elif filename.endswith('.css'):
            content_type = 'text/css'
        elif filename.endswith('.html'):
            content_type = 'text/html'
        elif filename.endswith('.json'):
            content_type = 'application/json'

        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        logger.info(f"Served static file: /static/{filename}")
        return Response(content, mimetype=content_type)

    except Exception as e:
        logger.error(f"Error serving static file {filename}: {e}")
        return f"Error: {e}", 500

# ルート直下のアセットもプレビューとして処理(CSS未適用問題の対策)
@app.route("/<path:filename>")
def catch_all_assets(filename):
    """
    /preview/ 配下でない直接リクエスト(/styles.css等)を
    プレビューエンドポイントに転送
    """
    # 既存のAPIルートと競合しないようにチェック
    excluded_paths = ['api', 'health', 'chat-message', 'favicon.ico', 'static']

    # パスの最初の部分をチェック
    first_segment = filename.split('/')[0]
    if first_segment in excluded_paths:
        return jsonify({"error": "Not found"}), 404

    # 静的アセットの拡張子チェック
    asset_extensions = ['.css', '.js', '.png', '.jpg', '.jpeg', '.svg', '.ico',
                       '.woff', '.woff2', '.ttf', '.json', '.webp', '.gif']

    if any(filename.endswith(ext) for ext in asset_extensions):
        logger.info(f"Asset request redirected: /{filename} -> /preview/{filename}")
        # プレビューエンドポイントに内部転送
        return serve_preview(filename)

    return jsonify({"error": "Not found"}), 404

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    logger.info(f"Starting HP Support Service on port {port}")
    logger.info(f"Prompts bucket: {os.getenv('PROMPTS_BUCKET_NAME', 'Not set - using defaults')}")
    logger.info(f"Preview bucket: {GCS_OUTPUT_BUCKET}/{GCS_OUTPUT_PATH}")
    logger.info(f"Default preview URL: {DEFAULT_PREVIEW_URL}")
    # ★★★ 追加部分 5: 起動時にTTSの認証情報を確認するログを追加 ★★★
    logger.info(f"Google Cloud TTS will be initialized on first request if credentials are set.")
    app.run(debug=False, host="0.0.0.0", port=port)
